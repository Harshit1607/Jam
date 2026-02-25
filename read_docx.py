import docx

doc = docx.Document(r"e:\Jam\Jam\IA - Jam the web.docx")
with open(r"e:\Jam\Jam\docx_output_utf8.txt", 'w', encoding='utf-8') as f:
    f.write("--- PARAGRAPHS ---\n")
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + "\n")

    f.write("--- TABLES ---\n")
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            f.write(" | ".join(row_data) + "\n")
